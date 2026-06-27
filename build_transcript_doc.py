#!/usr/bin/env python3
import argparse
import re
from datetime import datetime
from pathlib import Path

import av
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TIME_RE = re.compile(r"^\[(?P<start>\d\d:\d\d:\d\d) - (?P<end>\d\d:\d\d:\d\d)\]\s*(?P<text>.*)$")

BLACK = RGBColor(0, 0, 0)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_GRAY = "F2F4F7"


def audio_duration(audio_path: Path) -> float | None:
    try:
        with av.open(str(audio_path)) as container:
            if container.duration is not None:
                return float(container.duration / av.time_base)
            stream = container.streams.audio[0]
            if stream.duration is not None:
                return float(stream.duration * stream.time_base)
    except Exception:
        return None
    return None


def fmt_duration(seconds: float | None) -> str:
    if seconds is None:
        return "Unknown"
    total = int(round(seconds))
    h = total // 3600
    m = (total % 3600) // 60
    s = total % 60
    if h:
        return f"{h}h {m}m {s}s"
    if m:
        return f"{m}m {s}s"
    return f"{s}s"


def parse_transcript(transcript_path: Path) -> list[dict[str, str]]:
    entries = []
    for line in transcript_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        match = TIME_RE.match(line)
        if match:
            entries.append(match.groupdict())
        elif entries:
            entries[-1]["text"] = f"{entries[-1]['text']} {line}".strip()
        else:
            entries.append({"start": "", "end": "", "text": line})
    return entries


def set_run_font(run, size=11, bold=None, color=BLACK, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def para_spacing(paragraph, before=0, after=6, line=1.12):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.12

    header = section.header.paragraphs[0]
    para_spacing(header, after=0)
    set_run_font(header.add_run("Audio transcription document"), size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_spacing(footer, after=0)
    set_run_font(footer.add_run("Generated locally from automatic speech recognition; review key facts against the original audio."), size=8.5, color=MUTED)


def add_metadata_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1500, 8300])
    for row, (label, value) in zip(table.rows, rows):
        shade_cell(row.cells[0], LIGHT_GRAY)
        for cell, text, bold in [(row.cells[0], label, True), (row.cells[1], value, False)]:
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            para_spacing(paragraph, after=2)
            set_run_font(paragraph.add_run(text), size=10.5, bold=bold)


def build_markdown(title: str, audio_path: Path, transcript_path: Path, entries: list[dict[str, str]], output_path: Path):
    duration = fmt_duration(audio_duration(audio_path))
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    lines = [
        f"# {title}",
        "",
        "## Metadata",
        "",
        f"- Audio file: `{audio_path}`",
        f"- Duration: {duration}",
        f"- Generated at: {generated_at}",
        f"- Transcript source: `{transcript_path}`",
        "- Note: Automatic transcription may contain omissions or recognition errors. Review important facts against the original audio.",
        "",
        "## Transcript",
        "",
    ]
    for entry in entries:
        if entry["start"]:
            lines.append(f"### {entry['start']} - {entry['end']}")
        lines.append(entry["text"])
        lines.append("")
    output_path.write_text("\n".join(lines), encoding="utf-8")


def build_docx(title: str, audio_path: Path, transcript_path: Path, entries: list[dict[str, str]], output_path: Path):
    doc = Document()
    configure_doc(doc)

    p = doc.add_paragraph()
    para_spacing(p, before=10, after=4)
    set_run_font(p.add_run(title), size=22, bold=True)

    p = doc.add_paragraph()
    para_spacing(p, after=14)
    set_run_font(p.add_run("音频转录文档"), size=13, color=MUTED)

    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    rows = [
        ("音频文件", str(audio_path)),
        ("音频时长", fmt_duration(audio_duration(audio_path))),
        ("生成时间", generated_at),
        ("转录来源", str(transcript_path)),
        ("说明", "本文件由本地自动语音识别生成，可能存在错字、漏字和说话人混淆；关键事实请结合原录音复核。"),
    ]
    add_metadata_table(doc, rows)

    p = doc.add_paragraph()
    para_spacing(p, before=16, after=8)
    set_run_font(p.add_run("转录正文"), size=16, bold=True, color=BLUE)

    for entry in entries:
        p = doc.add_paragraph()
        para_spacing(p, before=4, after=2)
        if entry["start"]:
            set_run_font(p.add_run(f"{entry['start']} - {entry['end']}"), size=10.5, bold=True, color=DARK_BLUE)
            p.add_run("\n")
        set_run_font(p.add_run(entry["text"]), size=11, color=BLACK)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Build Markdown and Word documents from a timestamped transcript.")
    parser.add_argument("audio")
    parser.add_argument("--transcript", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--title", default=None)
    args = parser.parse_args()

    audio_path = Path(args.audio).expanduser().resolve()
    transcript_path = Path(args.transcript).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()
    title = args.title or f"{audio_path.stem} 转录文档"

    entries = parse_transcript(transcript_path)
    if not entries:
        raise RuntimeError(f"No transcript entries found in {transcript_path}")

    output_dir.mkdir(parents=True, exist_ok=True)
    md_path = output_dir / f"{audio_path.stem}_转录文档.md"
    docx_path = output_dir / f"{audio_path.stem}_转录文档.docx"
    build_markdown(title, audio_path, transcript_path, entries, md_path)
    build_docx(title, audio_path, transcript_path, entries, docx_path)

    print(f"Saved Markdown: {md_path}")
    print(f"Saved Word document: {docx_path}")


if __name__ == "__main__":
    main()
