#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import textwrap
import urllib.error
import urllib.request
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DEFAULT_OLLAMA_URL = "http://127.0.0.1:11434"
TIME_RE = re.compile(r"^\[(?P<start>\d\d:\d\d:\d\d) - (?P<end>\d\d:\d\d:\d\d)\]\s*(?P<text>.*)$")

BLACK = RGBColor(0, 0, 0)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(90, 90, 90)


def normalize_url(url: str) -> str:
    return url.rstrip("/")


def request_json(url: str, payload: dict | None = None, timeout: int = 600) -> dict:
    data = None
    headers = {}
    if payload is not None:
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        headers["Content-Type"] = "application/json"
    req = urllib.request.Request(url, data=data, headers=headers)
    try:
        with urllib.request.urlopen(req, timeout=timeout) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.URLError as exc:
        raise RuntimeError(f"无法连接 Ollama：{exc}") from exc
    except json.JSONDecodeError as exc:
        raise RuntimeError("Ollama 返回内容不是有效 JSON。") from exc


def request_json_stream(url: str, payload: dict, timeout: int = 900):
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(url, data=data, headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=timeout) as response:
            for raw_line in response:
                line = raw_line.decode("utf-8").strip()
                if not line:
                    continue
                try:
                    yield json.loads(line)
                except json.JSONDecodeError as exc:
                    raise RuntimeError("Ollama 流式返回内容不是有效 JSON。") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"无法连接 Ollama：{exc}") from exc


def list_models(ollama_url: str) -> list[str]:
    data = request_json(f"{normalize_url(ollama_url)}/api/tags", timeout=20)
    models = data.get("models", [])
    return [model.get("name", "") for model in models if model.get("name")]


def choose_model(ollama_url: str, preferred: str | None = None) -> str:
    if preferred:
        return preferred
    models = list_models(ollama_url)
    if not models:
        raise RuntimeError("Ollama 当前没有可用模型。请先在终端运行：ollama pull qwen2.5:7b")
    priority = ("qwen", "deepseek", "glm", "yi", "llama", "mistral")
    for key in priority:
        for model in models:
            if key in model.lower():
                return model
    return models[0]


def ollama_generate(ollama_url: str, model: str, prompt: str, temperature: float, num_ctx: int, num_predict: int) -> str:
    payload = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "你是严谨的中文会议记录助手。只输出用户要求的最终内容，不输出推理过程。不要使用 LaTeX、数学公式或 $...$ 包裹符号。",
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
        "stream": True,
        "think": False,
        "options": {
            "temperature": temperature,
            "num_ctx": num_ctx,
            "num_predict": num_predict,
        },
    }
    parts = []
    reported_chars = 0
    for data in request_json_stream(f"{normalize_url(ollama_url)}/api/chat", payload=payload, timeout=900):
        if data.get("error"):
            raise RuntimeError(f"Ollama 生成失败：{data['error']}")
        content = data.get("message", {}).get("content", "")
        if content:
            parts.append(content)
            current_chars = sum(len(part) for part in parts)
            if current_chars - reported_chars >= 120:
                reported_chars = current_chars
                print(f"Ollama generated {current_chars} characters...", flush=True)
    response = "".join(parts).strip()
    if not response:
        raise RuntimeError(f"Ollama 没有返回有效内容。请尝试更换模型或增大 num_ctx。模型：{model}")
    return response


def read_transcript(transcript_path: Path) -> str:
    lines = []
    for raw in transcript_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        match = TIME_RE.match(line)
        if match:
            text = match.group("text").strip()
            if text:
                lines.append(f"[{match.group('start')} - {match.group('end')}] {text}")
        else:
            lines.append(line)
    if not lines:
        raise RuntimeError(f"转写文本为空：{transcript_path}")
    return "\n".join(lines)


def split_text(text: str, max_chars: int) -> list[str]:
    if len(text) <= max_chars:
        return [text]

    chunks = []
    current = []
    current_len = 0
    for line in text.splitlines():
        line_len = len(line) + 1
        if current and current_len + line_len > max_chars:
            chunks.append("\n".join(current))
            current = []
            current_len = 0
        current.append(line)
        current_len += line_len
    if current:
        chunks.append("\n".join(current))
    return chunks


def chunk_prompt(chunk: str, idx: int, total: int) -> str:
    return textwrap.dedent(
        f"""
        你是严谨的中文会议记录助手。下面是一段录音自动转写文本，这是第 {idx}/{total} 段。
        请只根据转写内容整理，不要编造没有出现的信息。转写可能有错字、漏字，请把不确定处标为“待核实”。
        不要输出推理过程或思考过程，只输出整理结果。

        输出要求：
        1. 用中文。
        2. 这是给最终会议纪要使用的“详尽分段材料”，不要过度压缩。
        3. 保留重要事实、冲突点、诉求、回应、承诺、待办、金额/日期/姓名/机构/地点等具体信息。
        4. 按时间顺序整理关键进展；每个要点尽量带时间戳，方便回听。
        5. 对每个关键议题写清楚：谁提出、提出了什么、对方如何回应、是否形成结论或待办。
        6. 如果同一问题反复讨论，要合并归纳，但保留关键变化和反复点。
        7. 输出不少于 800 字；内容不足时也不要编造，改为更细地拆解已有内容。

        请使用以下结构：

        ## 本段关键事实
        - ...

        ## 本段时间线
        - [时间戳] ...

        ## 本段议题与观点
        ### 议题一：...
        - 提出/诉求：...
        - 回应/解释：...
        - 分歧/风险：...
        - 暂定结论/后续动作：...

        ## 本段待核实事项
        - ...

        转写文本：
        {chunk}
        """
    ).strip()


def final_prompt(source: str, title: str) -> str:
    return textwrap.dedent(
        f"""
        你是严谨的中文会议纪要整理助手。请基于以下录音转写或分段摘要，生成一份“详细、正式、可追溯”的会议材料。
        重要规则：
        - 只能根据输入内容整理，不要补充未出现的事实。
        - 自动转写可能有识别错误；遇到含糊、矛盾或无法确认的信息，请写入“待核实事项”。
        - 尽量保留关键时间戳，便于回听。
        - 重点突出参与各方诉求、争议点、回应、处理方案、承诺、后续动作。
        - 语言客观中立，避免情绪化判断。
        - 不要写得太简略。除非原文非常短，否则总字数建议不少于 2500 字。
        - 每个核心议题都要展开写：背景、事实依据、各方观点、争议焦点、已有回应、未解决问题。
        - 对结论和行动项要写清责任方、动作、时间要求；如果原文没有说明，请标“未明确”。
        - 可以对自动转写中的口语、重复、错字做适度整理，但不能改变事实含义。
        - 不要使用 LaTeX 或数学公式写法；箭头请直接写“→”，不要写“$\\rightarrow$”。
        - 不要输出推理过程或思考过程，只输出最终纪要。

        请严格使用以下 Markdown 结构：

        # {title}

        ## 一、核心摘要
        用 5-8 条概括本次沟通最重要的事实、争议、结论和风险。每条尽量包含时间戳或可回听线索。
        - ...

        ## 二、详细背景
        说明会议/沟通发生的背景、涉及事项、各方关注点，以及需要解决的问题。
        ...

        ## 三、按时间顺序的详细经过
        - [时间戳] 详细描述该阶段发生了什么、谁表达了什么、对后续讨论有什么影响。
        - [时间戳] ...

        ## 四、核心议题展开
        ### 议题一：...
        - 相关时间戳：...
        - 事实与背景：...
        - 主要诉求：...
        - 回应与解释：...
        - 分歧与风险：...
        - 当前状态：...

        ### 议题二：...
        - 相关时间戳：...
        - 事实与背景：...
        - 主要诉求：...
        - 回应与解释：...
        - 分歧与风险：...
        - 当前状态：...

        ## 五、各方观点与诉求
        如果转写中无法准确判断真实身份，请用“发言方/对方/相关方”等中性称呼。
        - 发言方 A：...
        - 发言方 B：...

        ## 六、已形成的结论或共识
        - ...

        ## 七、后续行动清单
        - 责任方：...
          事项：...
          时间要求：...
          备注：...

        ## 八、待核实事项
        - ...

        ## 九、风险与注意事项
        - ...

        ## 原始转写风险提示
        - 本纪要由本地 Ollama 大模型根据自动转写生成，关键事实仍需结合原录音复核。

        输入内容：
        {source}
        """
    ).strip()


def set_run_font(run, size=11, bold=None, color=BLACK):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def para_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def normalize_generated_markdown(text: str) -> str:
    replacements = {
        r"$\rightarrow$": "→",
        r"$\\rightarrow$": "→",
        r"\rightarrow": "→",
        r"\\rightarrow": "→",
        r"$\Rightarrow$": "⇒",
        r"$\\Rightarrow$": "⇒",
        r"\Rightarrow": "⇒",
        r"\\Rightarrow": "⇒",
        r"$\leftarrow$": "←",
        r"$\\leftarrow$": "←",
        r"\leftarrow": "←",
        r"\\leftarrow": "←",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"\$([→⇒←])\$", r"\1", text)
    return text


def add_inline_markdown_runs(paragraph, text: str, size=11, color=BLACK):
    text = normalize_generated_markdown(text)
    pos = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > pos:
            set_run_font(paragraph.add_run(text[pos : match.start()]), size=size, color=color)
        set_run_font(paragraph.add_run(match.group(1)), size=size, bold=True, color=color)
        pos = match.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]), size=size, color=color)


def add_markdown_line(doc: Document, line: str):
    stripped = normalize_generated_markdown(line.strip())
    if not stripped:
        return

    if stripped.startswith("# "):
        p = doc.add_paragraph()
        para_spacing(p, before=10, after=6)
        set_run_font(p.add_run(stripped[2:].strip()), size=20, bold=True)
        return

    if stripped.startswith("## "):
        p = doc.add_paragraph()
        para_spacing(p, before=14, after=6)
        set_run_font(p.add_run(stripped[3:].strip()), size=15, bold=True, color=BLUE)
        return

    if stripped.startswith("### "):
        p = doc.add_paragraph()
        para_spacing(p, before=8, after=4)
        set_run_font(p.add_run(stripped[4:].strip()), size=12.5, bold=True)
        return

    bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
    if bullet_match:
        p = doc.add_paragraph(style="List Bullet")
        para_spacing(p, after=4)
        add_inline_markdown_runs(p, bullet_match.group(1), size=11)
        return

    numbered_match = re.match(r"^\d+[.、]\s+(.*)$", stripped)
    if numbered_match:
        p = doc.add_paragraph(style="List Number")
        para_spacing(p, after=4)
        add_inline_markdown_runs(p, numbered_match.group(1), size=11)
        return

    p = doc.add_paragraph()
    para_spacing(p, after=5)
    add_inline_markdown_runs(p, stripped, size=11)


def build_docx(markdown_text: str, transcript_path: Path, model: str, output_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.add_run("Generated locally with Ollama; review key facts against the original audio."), size=8.5, color=MUTED)

    for line in markdown_text.splitlines():
        add_markdown_line(doc, line)

    p = doc.add_paragraph()
    para_spacing(p, before=12, after=4)
    set_run_font(p.add_run("生成信息"), size=12.5, bold=True, color=BLUE)
    for label, value in [
        ("Ollama 模型", model),
        ("转写文件", str(transcript_path)),
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(f"{label}：{value}"), size=10, color=MUTED)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate an outline and meeting minutes from a transcript with local Ollama.")
    parser.add_argument("--transcript", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--title", default="录音重点大纲与会议纪要")
    parser.add_argument("--ollama-url", default=DEFAULT_OLLAMA_URL)
    parser.add_argument("--model", default=None)
    parser.add_argument("--chunk-chars", type=int, default=5500)
    parser.add_argument("--num-ctx", type=int, default=12288)
    parser.add_argument("--num-predict", type=int, default=4200)
    parser.add_argument("--temperature", type=float, default=0.2)
    args = parser.parse_args()

    transcript_path = Path(args.transcript).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    model = choose_model(args.ollama_url, args.model.strip() if args.model else None)
    transcript = read_transcript(transcript_path)
    chunks = split_text(transcript, args.chunk_chars)

    print(f"Ollama URL: {normalize_url(args.ollama_url)}", flush=True)
    print(f"Ollama model: {model}", flush=True)
    print(f"Transcript chunks: {len(chunks)}", flush=True)

    if len(chunks) == 1:
        source = chunks[0]
    else:
        summaries = []
        for idx, chunk in enumerate(chunks, start=1):
            print(f"Summarizing chunk {idx}/{len(chunks)}...", flush=True)
            summaries.append(ollama_generate(args.ollama_url, model, chunk_prompt(chunk, idx, len(chunks)), args.temperature, args.num_ctx, args.num_predict))
        source = "\n\n".join(f"## 分段摘要 {idx}\n{summary}" for idx, summary in enumerate(summaries, start=1))

    print("Generating final outline and meeting minutes...", flush=True)
    markdown_text = normalize_generated_markdown(
        ollama_generate(args.ollama_url, model, final_prompt(source, args.title), args.temperature, args.num_ctx, args.num_predict)
    )

    md_path = output_dir / f"{transcript_path.stem}_重点大纲与会议纪要.md"
    docx_path = output_dir / f"{transcript_path.stem}_重点大纲与会议纪要.docx"
    md_path.write_text(markdown_text + "\n", encoding="utf-8")
    build_docx(markdown_text, transcript_path, model, docx_path)

    print(f"Saved Ollama Markdown: {md_path}", flush=True)
    print(f"Saved Ollama Word document: {docx_path}", flush=True)


if __name__ == "__main__":
    main()
